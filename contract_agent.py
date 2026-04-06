"""
ARA Travel — Contract Intelligence Agent v2.0
==============================================
يراقب SharePoint كل ساعة → يقرأ أي عقد → يستخرج كل التفاصيل → يدخل Supabase

يدعم جميع أنواع العقود:
  🏨 فنادق    — أسعار الغرف، المواسم، المواسم، Municipality + VAT، سياسة الإلغاء
  🚗 مواصلات  — نقل مطار، بين المدن، بالساعة، إيجار سيارات
  🎯 أنشطة    — بالشخص، بالمجموعة، باقات كاملة
  🗺️ DMC      — برامج متكاملة، عمرة، تجارب
  🏎️ إيجار   — يومي، أسبوعي، شهري، فئات السيارات

التشغيل:
  python contract_agent.py          ← يشتغل مرة واحدة (للـ GitHub Actions)
  python contract_agent.py --test   ← يختبر بدون حفظ
"""

import os, json, re, sys, hashlib, asyncio
from pathlib import Path
from datetime import datetime
import httpx
import pdfplumber
from supabase import create_client

# ── CONFIG ──────────────────────────────────────────────────────
ANTHROPIC_KEY  = os.environ.get("ANTHROPIC_API_KEY", "")
SUPABASE_URL   = os.environ.get("SUPABASE_URL", "")
SUPABASE_KEY   = os.environ.get("SUPABASE_KEY", "")
TENANT_ID      = os.environ.get("TENANT_ID", "")
CLIENT_ID      = os.environ.get("CLIENT_ID", "")
CLIENT_SECRET  = os.environ.get("CLIENT_SECRET", "")

# SharePoint — المجلد الجديد للعقود
SP_SITE        = "SourcingandContractingDepartment"
SP_CONTRACTS_FOLDER = os.environ.get(
    "SP_CONTRACTS_FOLDER",
    "/sites/SourcingandContractingDepartment/Shared Documents/ARA Brain Contracts"
)

MODEL          = "claude-sonnet-4-20250514"
TEST_MODE      = "--test" in sys.argv or os.environ.get("TEST_MODE") == "true"

supa = create_client(SUPABASE_URL, SUPABASE_KEY) if SUPABASE_URL and SUPABASE_KEY else None


# ══════════════════════════════════════════════════════════════════
# STEP 1: SHAREPOINT — اقرأ الملفات الجديدة
# ══════════════════════════════════════════════════════════════════

async def get_graph_token() -> str:
    async with httpx.AsyncClient() as http:
        r = await http.post(
            f"https://login.microsoftonline.com/{TENANT_ID}/oauth2/v2.0/token",
            data={
                "grant_type":    "client_credentials",
                "client_id":     CLIENT_ID,
                "client_secret": CLIENT_SECRET,
                "scope":         "https://graph.microsoft.com/.default"
            }
        )
        data = r.json()
        if "access_token" not in data:
            raise Exception(f"Token error: {data}")
        return data["access_token"]


async def get_new_contracts(token: str) -> list:
    """
    يجلب الملفات من SharePoint ويتجاهل ما تم معالجته مسبقاً.
    يتحقق من الـ file_hash في Supabase لتجنب التكرار.
    """
    processed_hashes = set()
    if supa:
        try:
            resp = supa.table("contract_import_log") \
                .select("file_hash") \
                .execute()
            processed_hashes = {r["file_hash"] for r in (resp.data or [])}
        except:
            pass  # الجدول غير موجود بعد — سيُنشأ لاحقاً

    async with httpx.AsyncClient(timeout=30) as http:
        # جلب قائمة الملفات من المجلد
        r = await http.get(
            f"https://graph.microsoft.com/v1.0/sites/{SP_SITE}"
            f":/drive/root:{SP_CONTRACTS_FOLDER}:/children"
            "?$select=id,name,size,lastModifiedDateTime,@microsoft.graph.downloadUrl"
            "&$top=200",
            headers={"Authorization": f"Bearer {token}"}
        )
        
        if r.status_code != 200:
            print(f"  ⚠️  SharePoint error {r.status_code}: {r.text[:200]}")
            return []

        files = r.json().get("value", [])

    # فلتر: PDF و Excel و Word فقط، وغير مُعالَج مسبقاً
    new_files = []
    for f in files:
        name = f.get("name", "")
        ext  = Path(name).suffix.lower()
        if ext not in [".pdf", ".xlsx", ".xls", ".docx", ".doc"]:
            continue

        # hash بسيط من الاسم + الحجم + تاريخ التعديل
        file_sig = f"{name}_{f.get('size',0)}_{f.get('lastModifiedDateTime','')}"
        file_hash = hashlib.md5(file_sig.encode()).hexdigest()

        if file_hash not in processed_hashes:
            f["file_hash"] = file_hash
            new_files.append(f)

    print(f"  📁 SharePoint: {len(files)} ملف إجمالي | {len(new_files)} جديد")
    return new_files


async def download_file(token: str, file_info: dict) -> bytes:
    """تحميل ملف من SharePoint"""
    download_url = file_info.get("@microsoft.graph.downloadUrl", "")
    if not download_url:
        # جلب رابط التحميل
        async with httpx.AsyncClient(timeout=30) as http:
            r = await http.get(
                f"https://graph.microsoft.com/v1.0/sites/{SP_SITE}"
                f"/drive/items/{file_info['id']}/content",
                headers={"Authorization": f"Bearer {token}"},
                follow_redirects=True
            )
            return r.content

    async with httpx.AsyncClient(timeout=60) as http:
        r = await http.get(download_url)
        return r.content


# ══════════════════════════════════════════════════════════════════
# STEP 2: READ FILE — قراءة المحتوى بكل أنواعه
# ══════════════════════════════════════════════════════════════════

def read_pdf_content(file_bytes: bytes) -> dict:
    """قراءة PDF واستخراج النص والجداول"""
    import tempfile
    result = {"text": "", "tables": [], "pages": 0}

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        with pdfplumber.open(tmp_path) as pdf:
            result["pages"] = len(pdf.pages)
            page_texts = []

            for i, page in enumerate(pdf.pages):
                page_content = []

                # نص
                text = page.extract_text() or ""
                if text.strip():
                    page_content.append(f"[PAGE {i+1}]\n{text}")

                # جداول
                tables = page.extract_tables()
                for t_idx, table in enumerate(tables):
                    if not table: continue
                    rows = []
                    for row in table:
                        clean = [str(c or "").strip().replace("\n", " ") for c in row]
                        if any(c for c in clean):
                            rows.append(" | ".join(clean))
                    if rows:
                        table_text = f"\n[TABLE {t_idx+1}]\n" + "\n".join(rows) + "\n"
                        page_content.append(table_text)
                        result["tables"].append(rows)

                page_texts.append("\n".join(page_content))

            result["text"] = "\n\n".join(page_texts)
    finally:
        Path(tmp_path).unlink(missing_ok=True)

    return result


def read_excel_content(file_bytes: bytes) -> dict:
    """قراءة Excel واستخراج البيانات"""
    import tempfile
    try:
        import openpyxl
    except ImportError:
        return {"text": "[Excel - openpyxl not installed]", "tables": [], "pages": 0}

    result = {"text": "", "tables": [], "pages": 0}

    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        wb = openpyxl.load_workbook(tmp_path, read_only=True, data_only=True)
        all_text = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            sheet_rows = []
            for row in ws.iter_rows(values_only=True):
                clean = [str(c or "").strip() for c in row]
                if any(c for c in clean):
                    sheet_rows.append(" | ".join(clean))

            if sheet_rows:
                all_text.append(f"[SHEET: {sheet_name}]\n" + "\n".join(sheet_rows[:200]))
                result["tables"].append(sheet_rows[:100])

        result["text"] = "\n\n".join(all_text)
        result["pages"] = len(wb.sheetnames)
    finally:
        Path(tmp_path).unlink(missing_ok=True)

    return result


def read_word_content(file_bytes: bytes) -> dict:
    """قراءة Word واستخراج النص"""
    import tempfile
    try:
        import docx
    except ImportError:
        return {"text": "[Word - python-docx not installed]", "tables": [], "pages": 1}

    result = {"text": "", "tables": [], "pages": 1}

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        doc = docx.Document(tmp_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        table_texts = []
        for t_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    rows.append(" | ".join(cells))
            if rows:
                table_texts.append(f"[TABLE {t_idx+1}]\n" + "\n".join(rows))
                result["tables"].append(rows)

        result["text"] = "\n".join(paragraphs) + "\n\n" + "\n\n".join(table_texts)
    finally:
        Path(tmp_path).unlink(missing_ok=True)

    return result


def read_file(file_bytes: bytes, filename: str) -> dict:
    """يقرأ الملف حسب نوعه"""
    ext = Path(filename).suffix.lower()
    print(f"  📖 قراءة {filename} ({len(file_bytes):,} bytes)...")

    if ext == ".pdf":
        return read_pdf_content(file_bytes)
    elif ext in [".xlsx", ".xls"]:
        return read_excel_content(file_bytes)
    elif ext in [".docx", ".doc"]:
        return read_word_content(file_bytes)
    else:
        return {"text": file_bytes.decode("utf-8", errors="ignore"), "tables": [], "pages": 1}


# ══════════════════════════════════════════════════════════════════
# STEP 3: AI EXTRACTION — استخراج البيانات بدقة كاملة
# ══════════════════════════════════════════════════════════════════

# البرومت الرئيسي — يفهم كل أنواع العقود
MASTER_EXTRACTION_PROMPT = """أنت محلل عقود متخصص لشركة ARA Travel، DMC سعودية.

مهمتك: قراءة أي عقد واستخراج كل التفاصيل التجارية بدقة 100% في JSON.

⚠️ قواعد صارمة:
1. استخرج الأسعار كما هي في العقد — لا تغير ولا تقرّب
2. إذا السعر غير واضح → net_price = null وضع في review_reasons
3. إذا شامل VAT → includes_vat = true
4. Municipality fee (رسوم البلدية) مختلفة عن VAT — وضّحها
5. أعِد JSON فقط — لا شرح ولا markdown

═══════════════════════════════════
JSON المطلوب:
═══════════════════════════════════

{
  "supplier_name": "اسم المورد من العقد",
  "supplier_name_ar": "الاسم العربي إن وجد",
  "contract_type": "hotel | transport_transfer | transport_rental | activity | dmc | restaurant | other",
  "contract_number": "رقم العقد أو null",
  "destination_city": "المدينة: Jeddah | Riyadh | AlUla | Makkah | Medina | Tabuk | Aseer | Hail | Taif | Other",
  "destination_country": "SA",
  "currency": "SAR | USD | EUR",
  "valid_from": "YYYY-MM-DD أو null",
  "valid_to": "YYYY-MM-DD أو null",
  
  "vat_info": {
    "vat_rate": 15.0,
    "municipality_fee_rate": 0.0,
    "total_tax_rate": 15.0,
    "prices_include_tax": true,
    "notes": "شرح نظام الضريبة في هذا العقد"
  },
  
  "cancellation_policy": {
    "low_season_days": 3,
    "high_season_days": 7,
    "peak_season_days": 14,
    "no_show_policy": "نص السياسة",
    "notes": "أي شروط خاصة"
  },
  
  "payment_terms": "شروط الدفع كما في العقد",
  "special_conditions": ["قائمة الشروط المهمة"],
  
  "rate_lines": [
    {
      "service_name": "اسم الخدمة بالضبط من العقد",
      "service_name_ar": "الاسم العربي",
      "category": "room | suite | villa | airport_transfer | intercity | hourly | daily_rental | activity_person | activity_group | package | meal | other",
      
      "net_price": 1000.00,
      "net_price_child": null,
      "net_price_infant": 0,
      
      "includes_tax": true,
      "vat_included": true,
      "municipality_included": false,
      
      "unit": "per_night | per_room | per_person | per_vehicle | per_group | per_trip | per_day | per_week | per_month | per_hour",
      "max_pax": null,
      "min_pax": 1,
      
      "room_type": "standard | superior | deluxe | suite | villa | null",
      "meal_plan": "RO | BB | HB | FB | AI | null",
      "hotel_stars": null,
      
      "vehicle_type": "sedan | suv | van | minibus | bus | limo | boat | null",
      "vehicle_model": "Toyota Camry | Chevrolet Tahoe | etc أو null",
      
      "season": "low | shoulder | high | peak | year_round | null",
      "valid_from": "YYYY-MM-DD أو null",
      "valid_to": "YYYY-MM-DD أو null",
      "season_dates": "وصف تواريخ الموسم",
      
      "route_from": "نقطة الانطلاق للمواصلات",
      "route_to": "الوجهة للمواصلات",
      
      "duration_hours": null,
      "duration_minutes": null,
      "min_nights": 1,
      "advance_booking_days": null,
      "release_days": null,
      
      "includes": ["ما يشمله السعر"],
      "excludes": ["ما لا يشمله السعر"],
      
      "notes": "أي ملاحظة مهمة",
      "original_text": "النص الحرفي من العقد (أقل من 200 حرف)"
    }
  ],
  
  "supplements": [
    {
      "name": "اسم الإضافة",
      "price": 0.0,
      "unit": "per_person | per_night | per_room",
      "notes": ""
    }
  ],
  
  "ai_confidence": 0.95,
  "needs_human_review": false,
  "review_reasons": ["أسباب إذا احتاج مراجعة"],
  "extraction_notes": "ملاحظات مهمة للفريق"
}

═══════════════════════════════════
إرشادات حسب نوع العقد:
═══════════════════════════════════

🏨 الفنادق:
- استخرج كل موسم كـ rate_line منفصل (Low/Shoulder/High/Peak)
- Municipality fee شائعة في جدة ومكة — نسبتها عادة 5%
- سعر AI (All Inclusive) غير سعر BB (Breakfast Only)
- الـ supplement (extra adult, child) = rate_line منفصل

🚗 المواصلات (نقل):
- كل مسار + كل نوع سيارة = rate_line منفصل
- وضّح: هل السعر one-way أم return
- لا يشمل Umrah/Hajj season عادة

🏎️ الإيجار (Rental):
- يومي/أسبوعي/شهري = rate lines منفصلة
- الـ km المشمولة مهمة — سجّلها في includes
- CDW (تأمين) قد يكون منفصل — سجّله في supplements

🎯 الأنشطة:
- per_person vs per_group — فرّق بينهم
- المدة مهمة (duration_hours)
- الحد الأقصى للأشخاص (max_pax)

📦 الـ DMC (باقات):
- الباقة الكاملة = rate_line واحد
- وضّح ماذا يشمل وماذا لا يشمل بالتفصيل"""


async def extract_contract_data(file_content: dict, filename: str) -> dict:
    """يرسل محتوى الملف لـ Claude ويستخرج البيانات"""

    text = file_content["text"]
    MAX_CHARS = 14000

    # تقطيع ذكي: نحافظ على البداية (اسم المورد، نوع العقد) والنهاية (الشروط)
    if len(text) > MAX_CHARS:
        keep_start = int(MAX_CHARS * 0.70)
        keep_end   = int(MAX_CHARS * 0.30)
        text = (
            text[:keep_start]
            + f"\n\n[... {len(text) - MAX_CHARS:,} حرف محذوف ...]\n\n"
            + text[-keep_end:]
        )

    user_message = f"""اسم الملف: {filename}
الصفحات: {file_content['pages']} | الجداول: {len(file_content['tables'])}

محتوى العقد:
{'='*60}
{text}
{'='*60}

استخرج كل البيانات التجارية بدقة كاملة."""

    async with httpx.AsyncClient(timeout=120) as http:
        r = await http.post(
            "https://api.anthropic.com/v1/messages",
            headers={
                "x-api-key":         ANTHROPIC_KEY,
                "anthropic-version": "2023-06-01",
                "content-type":      "application/json"
            },
            json={
                "model":      MODEL,
                "max_tokens": 8000,
                "system":     MASTER_EXTRACTION_PROMPT,
                "messages":   [{"role": "user", "content": user_message}]
            }
        )

    raw = r.json()
    if "error" in raw:
        raise Exception(f"Claude error: {raw['error']}")

    response_text = raw["content"][0]["text"]
    response_text = re.sub(r'^```json\s*', '', response_text.strip())
    response_text = re.sub(r'\s*```$', '',  response_text)
    response_text = response_text.strip()

    try:
        return json.loads(response_text)
    except json.JSONDecodeError as e:
        # محاولة إصلاح
        match = re.search(r'\{.*\}', response_text, re.DOTALL)
        if match:
            try:
                return json.loads(match.group())
            except:
                pass
        raise Exception(f"JSON parse failed: {e}")


# ══════════════════════════════════════════════════════════════════
# STEP 4: SAVE TO SUPABASE — حفظ في قاعدة البيانات
# ══════════════════════════════════════════════════════════════════

# Markup الافتراضي لكل نوع خدمة
DEFAULT_MARKUP = {
    "hotel":              20.0,
    "transport_transfer": 15.0,
    "transport_rental":   15.0,
    "activity":           25.0,
    "dmc":                25.0,
    "restaurant":         30.0,
    "other":              25.0,
}

# تحويل نوع العقد لـ DB enum
CONTRACT_TYPE_MAP = {
    "hotel":              "hotel",
    "transport_transfer": "transport",
    "transport_rental":   "transport",
    "activity":           "activity_operator",
    "dmc":                "dmc",
    "restaurant":         "restaurant",
    "other":              "misc",
}


def save_to_supabase(extracted: dict, filename: str, file_hash: str) -> dict:
    """يحفظ العقد وكل أسطر الأسعار في Supabase"""

    result = {
        "success":         False,
        "contract_id":     None,
        "lines_saved":     0,
        "lines_skipped":   0,
        "errors":          [],
        "supplier":        extracted.get("supplier_name", "Unknown"),
    }

    if not supa:
        result["errors"].append("Supabase not connected")
        return result

    try:
        # ── GET/CREATE SUPPLIER ──────────────────────────────────
        supplier_name = extracted.get("supplier_name", "Unknown")[:200]
        c_type_raw    = extracted.get("contract_type", "other")
        c_type_db     = CONTRACT_TYPE_MAP.get(c_type_raw, "misc")
        markup_default = DEFAULT_MARKUP.get(c_type_raw, 20.0)

        # ابحث عن مورد موجود
        s_resp = supa.table("suppliers") \
            .select("id") \
            .ilike("name_en", f"%{supplier_name[:30]}%") \
            .limit(1).execute()

        if s_resp.data:
            supplier_id = s_resp.data[0]["id"]
        else:
            s_new = supa.table("suppliers").insert({
                "name_en":       supplier_name,
                "name_ar":       extracted.get("supplier_name_ar", supplier_name),
                "supplier_type": c_type_db if c_type_db != "misc" else "misc",
                "is_active":     True,
            }).execute()
            supplier_id = s_new.data[0]["id"] if s_new.data else None

        # ── GET DESTINATION ──────────────────────────────────────
        dest_city = extracted.get("destination_city", "")
        dest_id   = None
        if dest_city and dest_city != "Other":
            d_resp = supa.table("destinations") \
                .select("id") \
                .ilike("name_en", f"%{dest_city[:15]}%") \
                .limit(1).execute()
            if d_resp.data:
                dest_id = d_resp.data[0]["id"]

        # ── VAT INFO ─────────────────────────────────────────────
        vat_info     = extracted.get("vat_info", {})
        vat_rate     = float(vat_info.get("vat_rate", 15.0) or 15.0)
        muni_rate    = float(vat_info.get("municipality_fee_rate", 0.0) or 0.0)
        total_tax    = float(vat_info.get("total_tax_rate", 15.0) or 15.0)
        prices_incl  = bool(vat_info.get("prices_include_tax", False))

        # ── INSERT CONTRACT ───────────────────────────────────────
        cancel = extracted.get("cancellation_policy", {})
        contract_data = {
            "supplier_id":          supplier_id,
            "supplier_name":        supplier_name,
            "contract_type":        c_type_db,
            "contract_number":      extracted.get("contract_number"),
            "destination_id":       dest_id,
            "valid_from":           extracted.get("valid_from"),
            "valid_to":             extracted.get("valid_to"),
            "currency":             extracted.get("currency", "SAR"),
            "payment_terms":        extracted.get("payment_terms"),
            "cancellation_policy":  json.dumps(cancel, ensure_ascii=False),
            "source_filename":      filename,
            "ai_confidence":        float(extracted.get("ai_confidence", 0.8)),
            "needs_human_review":   bool(extracted.get("needs_human_review", False)),
            "ai_extraction_notes":  extracted.get("extraction_notes"),
            "status":               "active",
            "ai_extracted_at":      datetime.utcnow().isoformat(),
        }

        c_resp = supa.table("contracts").insert(contract_data).execute()
        if not c_resp.data:
            raise Exception("Contract insert failed")

        contract_id         = c_resp.data[0]["id"]
        result["contract_id"] = contract_id

        # ── INSERT RATE LINES ─────────────────────────────────────
        all_lines = extracted.get("rate_lines", []) + [
            {
                "service_name": s.get("name", "Supplement"),
                "category":     "supplement",
                "net_price":    s.get("price"),
                "unit":         s.get("unit", "per_person"),
                "notes":        s.get("notes", ""),
                "original_text": s.get("name", ""),
            }
            for s in extracted.get("supplements", [])
            if s.get("price") and float(str(s.get("price", 0) or 0)) > 0
        ]

        for line in all_lines:
            try:
                # تحقق من السعر
                net_raw = line.get("net_price")
                if net_raw is None or str(net_raw).strip() in ["", "null", "None"]:
                    result["lines_skipped"] += 1
                    continue

                net_price = float(str(net_raw).replace(",", "").replace(" ", ""))
                if net_price <= 0:
                    result["lines_skipped"] += 1
                    continue

                # Child price
                child_price = None
                if line.get("net_price_child") is not None:
                    try:
                        cp = float(str(line["net_price_child"]).replace(",", ""))
                        if cp > 0: child_price = cp
                    except: pass

                # VAT handling
                line_includes_vat  = bool(line.get("includes_tax", prices_incl))
                line_includes_muni = bool(line.get("municipality_included", False))
                supplier_vat = (vat_rate + muni_rate) if (line_includes_vat and line_includes_muni) \
                               else vat_rate if line_includes_vat else 0.0

                # Route
                route = None
                if line.get("route_from") or line.get("route_to"):
                    route = f"{line.get('route_from','?')} → {line.get('route_to','?')}"

                # Rate type من الـ category
                cat = line.get("category", "other")
                rate_type_map = {
                    "room": "room_night", "suite": "room_night", "villa": "room_night",
                    "airport_transfer": "transfer_oneway", "intercity": "transfer_oneway",
                    "hourly": "transfer_hourly", "daily_rental": "car_rental_daily",
                    "activity_person": "activity_person", "activity_group": "activity_group",
                    "package": "room_package", "supplement": "misc", "meal": "meal_person",
                }
                rate_type = rate_type_map.get(cat, "misc")

                # Original text
                orig_parts = [
                    line.get("original_text", ""),
                    route or "",
                    line.get("notes", ""),
                ]
                orig_text = " | ".join(p for p in orig_parts if p)[:500]

                rate_data = {
                    "contract_id":       contract_id,
                    "rate_line_type":    rate_type,
                    "service_name":      str(line.get("service_name", ""))[:200],
                    "service_name_ar":   str(line.get("service_name_ar", "") or "")[:200] or None,
                    "description":       str(line.get("category", "") or "")[:300] or None,
                    "destination_id":    dest_id,
                    "room_type":         line.get("room_type"),
                    "meal_plan":         line.get("meal_plan"),
                    "vehicle_type":      line.get("vehicle_type"),
                    "max_pax":           line.get("max_pax"),
                    "min_pax":           int(line.get("min_pax", 1) or 1),
                    "currency":          extracted.get("currency", "SAR"),
                    "net_price":         net_price,
                    "net_price_child":   child_price,
                    "includes_vat":      line_includes_vat,
                    "supplier_vat_rate": supplier_vat,
                    "markup_pct":        markup_default,
                    "our_vat_rate":      15.00,
                    "valid_from":        line.get("valid_from") or extracted.get("valid_from"),
                    "valid_to":          line.get("valid_to")   or extracted.get("valid_to"),
                    "season_type":       line.get("season"),
                    "min_nights":        int(line.get("min_nights", 1) or 1),
                    "release_days":      line.get("release_days"),
                    "route_name_ar":     route,
                    "original_text":     orig_text,
                    "is_active":         True,
                }

                supa.table("contract_rate_lines").insert(rate_data).execute()
                result["lines_saved"] += 1

            except Exception as e:
                result["errors"].append(f"Line error: {str(e)[:100]}")
                result["lines_skipped"] += 1

        # ── LOG PROCESSED FILE ────────────────────────────────────
        try:
            supa.table("contract_import_log").insert({
                "file_hash":    file_hash,
                "filename":     filename,
                "contract_id":  contract_id,
                "supplier":     supplier_name,
                "lines_saved":  result["lines_saved"],
                "imported_at":  datetime.utcnow().isoformat(),
            }).execute()
        except:
            pass  # الجدول سيُنشأ في الـ SQL setup

        result["success"] = result["lines_saved"] > 0

    except Exception as e:
        result["errors"].append(f"Fatal: {str(e)}")

    return result


# ══════════════════════════════════════════════════════════════════
# STEP 5: NOTIFY — إشعار الموظف
# ══════════════════════════════════════════════════════════════════

async def notify_staff(results_summary: list):
    """يرسل ملخص للموظف عبر واتساب أو يطبعه"""

    success = [r for r in results_summary if r.get("success")]
    failed  = [r for r in results_summary if not r.get("success")]
    review  = [r for r in success if r.get("needs_review")]

    message_lines = [
        f"✅ *Contract Agent — تقرير الاستيراد*",
        f"📅 {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        "",
        f"✅ تم استيراد: {len(success)} عقد",
        f"❌ فشل: {len(failed)} عقد",
        f"⚠️ تحتاج مراجعة: {len(review)} عقد",
        "",
    ]

    if success:
        message_lines.append("*العقود المستوردة:*")
        for r in success[:10]:
            message_lines.append(
                f"  • {r.get('supplier','?')} — {r.get('lines_saved',0)} سطر سعر"
            )

    if review:
        message_lines.append("\n*⚠️ تحتاج مراجعتك:*")
        for r in review:
            message_lines.append(f"  • {r.get('supplier','?')}")

    if failed:
        message_lines.append("\n*❌ فشل الاستيراد:*")
        for r in failed[:5]:
            message_lines.append(
                f"  • {r.get('filename','?')}: {r.get('errors',['?'])[0][:60]}"
            )

    print("\n" + "="*55)
    print("\n".join(message_lines))
    print("="*55 + "\n")


# ══════════════════════════════════════════════════════════════════
# MAIN PIPELINE
# ══════════════════════════════════════════════════════════════════

async def main():
    print(f"\n{'='*55}")
    print(f"📋 ARA Travel Contract Agent")
    print(f"   {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"   Mode: {'🔍 TEST' if TEST_MODE else '💾 LIVE'}")
    print(f"{'='*55}\n")

    results = []

    # ── وضع الاختبار: ارفع ملفات محلية ──────────────────────────
    if TEST_MODE:
        test_files = list(Path(".").glob("*.pdf")) + \
                     list(Path(".").glob("*.xlsx")) + \
                     list(Path(".").glob("*.docx"))

        if not test_files:
            print("⚠️  لا توجد ملفات محلية للاختبار")
            print("   ضع ملفات العقد في نفس المجلد وشغّل: python contract_agent.py --test")
            return

        print(f"🧪 وجد {len(test_files)} ملف للاختبار\n")

        for f in test_files:
            print(f"📄 {f.name}")
            file_bytes   = f.read_bytes()
            file_content = read_file(file_bytes, f.name)
            file_hash    = hashlib.md5(f.name.encode()).hexdigest()

            print(f"   {file_content['pages']} صفحات | {len(file_content['tables'])} جداول | {len(file_content['text']):,} حرف")

            try:
                extracted = await extract_contract_data(file_content, f.name)

                print(f"   ✅ المورد: {extracted.get('supplier_name')}")
                print(f"   ✅ النوع: {extracted.get('contract_type')}")
                print(f"   ✅ المدينة: {extracted.get('destination_city')}")
                print(f"   ✅ الأسطر: {len(extracted.get('rate_lines', []))}")
                print(f"   ✅ الثقة: {extracted.get('ai_confidence', 0):.0%}")

                if extracted.get("vat_info"):
                    v = extracted["vat_info"]
                    print(f"   💰 VAT: {v.get('vat_rate')}% | Municipality: {v.get('municipality_fee_rate')}% | شامل: {v.get('prices_include_tax')}")

                if extracted.get("needs_human_review"):
                    print(f"   ⚠️  يحتاج مراجعة: {extracted.get('review_reasons', [])}")

                # أول 3 أسطر سعر
                for line in extracted.get("rate_lines", [])[:3]:
                    print(f"   💵 {line.get('service_name','?')[:45]:<45} = {line.get('net_price','?')} SAR [{line.get('unit','?')}]")

                if not TEST_MODE:
                    db_result = save_to_supabase(extracted, f.name, file_hash)
                    print(f"   💾 حُفظ: {db_result['lines_saved']} سطر")
                    results.append({**db_result, "filename": f.name, "needs_review": extracted.get("needs_human_review")})
                else:
                    print(f"   🔍 TEST MODE — لم يُحفظ في DB")
                    results.append({"success": True, "supplier": extracted.get("supplier_name"), "lines_saved": len(extracted.get("rate_lines",[])), "filename": f.name})

            except Exception as e:
                print(f"   ❌ خطأ: {e}")
                results.append({"success": False, "filename": f.name, "errors": [str(e)]})

            print()

    # ── الوضع الحقيقي: من SharePoint ────────────────────────────
    else:
        if not TENANT_ID or not CLIENT_ID:
            print("⚠️  لا توجد Microsoft credentials — جرّب: python contract_agent.py --test")
            return

        print("🔑 جلب Graph Token...")
        token = await get_graph_token()
        print("✅ Token OK\n")

        print(f"📁 فحص SharePoint: {SP_CONTRACTS_FOLDER}")
        new_files = await get_new_contracts(token)

        if not new_files:
            print("✅ لا توجد عقود جديدة")
            return

        print(f"\n🆕 {len(new_files)} عقد جديد للمعالجة\n")

        for i, file_info in enumerate(new_files, 1):
            filename  = file_info["name"]
            file_hash = file_info["file_hash"]

            print(f"[{i}/{len(new_files)}] 📄 {filename}")

            try:
                # تحميل الملف
                file_bytes   = await download_file(token, file_info)
                file_content = read_file(file_bytes, filename)

                print(f"   {file_content['pages']} صفحات | {len(file_content['text']):,} حرف")

                # استخراج البيانات
                extracted = await extract_contract_data(file_content, filename)
                confidence = extracted.get("ai_confidence", 0)
                lines_count = len(extracted.get("rate_lines", []))

                print(f"   ✅ {extracted.get('supplier_name')} | {extracted.get('contract_type')} | {lines_count} سطر | {confidence:.0%}")

                # حفظ في Supabase
                db_result = save_to_supabase(extracted, filename, file_hash)

                print(f"   💾 حُفظ: {db_result['lines_saved']} سطر" + (f" | ⚠️ {db_result['lines_skipped']} تجاوز" if db_result['lines_skipped'] else ""))

                if db_result["errors"]:
                    for err in db_result["errors"][:2]:
                        print(f"   ❌ {err}")

                results.append({
                    **db_result,
                    "filename":     filename,
                    "needs_review": extracted.get("needs_human_review", False)
                })

            except Exception as e:
                print(f"   ❌ فشل: {e}")
                results.append({
                    "success":  False,
                    "filename": filename,
                    "supplier": filename,
                    "errors":   [str(e)]
                })

            print()

    # ── ملخص نهائي ───────────────────────────────────────────────
    if results:
        await notify_staff(results)


if __name__ == "__main__":
    asyncio.run(main())
